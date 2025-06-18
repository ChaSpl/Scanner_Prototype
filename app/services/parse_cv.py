# app/services/parse_cv.py

import os
import re
import json
from datetime import date
from docx import Document as DocxDocument
from dateutil import parser as date_parser
from db.session import SessionLocal
from app.utils.audit_logger import logger
from app.services.llm_cv_parser import parse_cv_with_llm
from app.models import (
    Document,
    Person,
    Education,
    Language,
    Experience,
    Certification,
    Award,
    FurtherEducation,
    Publication,
    PersonalAchievement,
    PrivateMilestone,
)
from app.services.generate_pdf import generate_cv_pdf
from app.services.plot_timeline_vertical import plot_timeline_and_save

def extract_text(docx_path: str) -> list[str]:
    doc = DocxDocument(docx_path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

def normalize_date(raw: str, prefer_start=True) -> tuple[date|None, str|None]:
    if not raw:
        return None, None
    raw = raw.strip().lower()
    if raw == "present":
        return None, None
    try:
        parsed = date_parser.parse(raw, fuzzy=True, default=date(1900,1,1))
        if re.fullmatch(r"\d{4}", raw):
            return date(int(raw),1,1), "year"
        if re.fullmatch(r"\d{4}[-/]\d{2}", raw):
            year, month = map(int, re.split("[-/]", raw))
            return date(year, month, 1), "month"
        return parsed.date(), "day"
    except Exception as e:
        logger.warning(f"⚠️ Date parse failed for '{raw}': {e}")
        return None, None

def normalize_proficiency(val: str) -> str|None:
    if not val:
        return None
    v = val.strip().lower()
    if v in ("mother tongue","native speaker","native"):
        return "native"
    if v in ("fluent","professional"):
        return "professional"
    if v in ("beginner","elementary","basic"):
        return "basic"
    return v

def get_or_create_document(session, cv_path: str, uploaded_by="system") -> Document:
    doc = session.query(Document).filter_by(source_filename=cv_path).first()
    if not doc:
        doc = Document(
            title=os.path.basename(cv_path),
            source_filename=cv_path,
            uploaded_by=uploaded_by,
            status="parsed"
        )
        session.add(doc)
        session.flush()
        logger.info(f"📄 Created new Document: {doc.title}")
    else:
        logger.info(f"📄 Found existing Document: {doc.title}")
    return doc

def get_or_create_person(
    session,
    data: dict,
    document: Document,
    fallback_email: str | None = None
) -> Person:
    """
    Given the parsed `data` and the Document record, find or create
    the Person. If `data["email"]` is missing, falls back to `fallback_email`.
    Raises ValueError if neither is provided.
    """
    # 1) Determine the email to use
    raw_email = data.get("email") or fallback_email
    email = (raw_email or "").strip().lower()
    if not email:
        raise ValueError("Email address is required to match a person.")

    # 2) Look up an existing Person
    person = session.query(Person).filter_by(email=email).first()
    if person:
        logger.info(f"🔁 Found existing person: {person.full_name}")
        # If this document isn’t already linked, attach it
        if person.document_id != document.id:
            person.document = document
            session.flush()
            logger.info(f"🔁 Linked person {person.full_name} to document {document.id}")
    else:
        # 3) Create a new Person record
        person = Person(
            full_name    = data.get("full_name", ""),
            email        = email,
            phone        = data.get("phone", ""),
            linkedin     = data.get("linkedin", ""),
            github       = data.get("github", ""),
            website      = data.get("website", ""),
            short_bio    = data.get("short_bio", ""),
            document     = document,
        )
        session.add(person)
        session.flush()
        logger.info(f"✅ Created new person: {person.full_name}")

    return person



########################################################################################################
#
# UPSERTS
#
#########################################################################################################


def upsert_educations(session, person, educations_data):
    existing = {
        (e.degree, e.field_of_study, e.institution): e
        for e in person.educations
    }

    for edu_data in educations_data:
        if not edu_data.get("degree") or not edu_data.get("institution"):
            continue   # skip incomplete entries

        key = (
            edu_data.get("degree"),
            edu_data.get("field"),
            edu_data.get("institution"),
        )
        start_date, start_precision = normalize_date(edu_data.get("start_date"))
        end_date, end_precision = normalize_date(edu_data.get("end_date"))

        existing_entry = existing.get(key)
        if existing_entry:
            updated = False
            if (
                existing_entry.start_date != start_date or
                existing_entry.end_date != end_date or
                existing_entry.start_date_precision != start_precision or
                existing_entry.end_date_precision != end_precision
            ):
                existing_entry.start_date = start_date
                existing_entry.start_date_precision = start_precision
                existing_entry.end_date = end_date
                existing_entry.end_date_precision = end_precision
                updated = True

            if updated:
                logger.info(f"🔁 Updated education | Person: {person.full_name} | Key: {key}")
        else:
            session.add(Education(
                person_id=person.id,
                degree=edu_data.get("degree"),
                field_of_study=edu_data.get("field"),
                institution=edu_data.get("institution"),
                start_date=start_date,
                start_date_precision=start_precision,
                end_date=end_date,
                end_date_precision=end_precision
            ))
            logger.info(f"➕ Added new education | Person: {person.full_name} | Key: {key}")



def upsert_experiences(session, person, experiences_data):
    existing = {
        (e.title, e.company, e.start_date): e
        for e in person.experiences
    }

    for exp_data in experiences_data:
        if not exp_data.get("title") or not exp_data.get("company"):
            continue

        start_date, start_precision = normalize_date(exp_data.get("start_date"))
        end_date, end_precision = normalize_date(exp_data.get("end_date"))
        key = (
            exp_data.get("title"),
            exp_data.get("company"),
            start_date
        )

        existing_entry = existing.get(key)
        if existing_entry:
            updated = False
            if (
                existing_entry.end_date != end_date or
                existing_entry.end_date_precision != end_precision or
                existing_entry.start_date_precision != start_precision or
                existing_entry.location != exp_data.get("location") or
                existing_entry.role_type != exp_data.get("role_type") or
                existing_entry.role_description != exp_data.get("role_description")
            ):
                existing_entry.end_date = end_date
                existing_entry.end_date_precision = end_precision
                existing_entry.start_date_precision = start_precision
                existing_entry.location = exp_data.get("location")
                existing_entry.role_type = exp_data.get("role_type")
                existing_entry.role_description = exp_data.get("role_description")
                updated = True

            if updated:
                logger.info(f"🔁 Updated experience | Person: {person.full_name} | Key: {key}")
        else:
            session.add(Experience(
                person_id=person.id,
                title=exp_data.get("title"),
                company=exp_data.get("company"),
                location=exp_data.get("location"),
                start_date=start_date,
                start_date_precision=start_precision,
                end_date=end_date,
                end_date_precision=end_precision,
                role_type=exp_data.get("role_type"),
                role_description=exp_data.get("role_description"),
            ))
            logger.info(f"➕ Added new experience | Person: {person.full_name} | Key: {key}")


def upsert_languages(session, person, languages_data):
    existing = {
        lang.language.lower(): lang
        for lang in person.languages
        if lang.language
    }

    for lang_data in languages_data:
        language = lang_data.get("language")
        if not language:
            continue

        norm_lang = language.strip().lower()
        written = normalize_proficiency(lang_data.get("proficiency_written"))
        spoken = normalize_proficiency(lang_data.get("proficiency_spoken"))

        existing_entry = existing.get(norm_lang)
        if existing_entry:
            updated = False
            if (
                existing_entry.proficiency_written != written or
                existing_entry.proficiency_spoken != spoken
            ):
                existing_entry.proficiency_written = written
                existing_entry.proficiency_spoken = spoken
                updated = True

            if updated:
                logger.info(f"🔁 Updated language | Person: {person.full_name}")
        else:
            session.add(Language(
                person_id=person.id,
                language=language,
                proficiency_written=written,
                proficiency_spoken=spoken
            ))
            logger.info(f"➕ Added new language | Person: {person.full_name}")



def upsert_further_education(session, person, fe_data_list):
    existing = {
        (fe.title, fe.institution): fe
        for fe in person.further_education
        if fe.title and fe.institution
    }

    for fe_data in fe_data_list:
        title = fe_data.get("title")
        institution = fe_data.get("institution")
        if not title or not institution:
            continue

        key = (title, institution)
        start_date, start_precision = normalize_date(fe_data.get("start_date"))
        end_date, end_precision = normalize_date(fe_data.get("end_date"))

        existing_entry = existing.get(key)
        if existing_entry:
            updated = False
            if (
                existing_entry.start_date != start_date or
                existing_entry.end_date != end_date or
                existing_entry.start_date_precision != start_precision or
                existing_entry.end_date_precision != end_precision
            ):
                existing_entry.start_date = start_date
                existing_entry.start_date_precision = start_precision
                existing_entry.end_date = end_date
                existing_entry.end_date_precision = end_precision
                updated = True

            if updated:
                logger.info(f"🔁 Updated further education | Person: {person.full_name} | Key: {key}")
        else:
            session.add(FurtherEducation(
                person_id=person.id,
                title=title,
                institution=institution,
                start_date=start_date,
                start_date_precision=start_precision,
                end_date=end_date,
                end_date_precision=end_precision
            ))
            logger.info(f"➕ Added new further education | Person: {person.full_name} | Key: {key}")


def upsert_certifications(session, person, certs_data):
    existing = {
        (c.name, c.issuer): c
        for c in person.certifications
        if c.name and c.issuer
    }

    for cert_data in certs_data:
        name = cert_data.get("name")
        issuer = cert_data.get("issuer")
        if not name or not issuer:
            continue

        key = (name, issuer)
        start_date, start_precision = normalize_date(cert_data.get("start_date"))
        end_date, end_precision = normalize_date(cert_data.get("end_date"))

        existing_entry = existing.get(key)
        if existing_entry:
            updated = False
            if (
                existing_entry.start_date != start_date or
                existing_entry.end_date != end_date or
                existing_entry.start_date_precision != start_precision or
                existing_entry.end_date_precision != end_precision
            ):
                existing_entry.start_date = start_date
                existing_entry.start_date_precision = start_precision
                existing_entry.end_date = end_date
                existing_entry.end_date_precision = end_precision
                updated = True

            if updated:
                logger.info(f"🔁 Updated certification | Person: {person.full_name} | Key: {key}")
        else:
            session.add(Certification(
                person_id=person.id,
                name=name,
                issuer=issuer,
                start_date=start_date,
                start_date_precision=start_precision,
                end_date=end_date,
                end_date_precision=end_precision
            ))
            logger.info(f"➕ Added new certification | Person: {person.full_name} | Key: {key}")


def upsert_awards(session, person, awards_data):
    existing = {
        (a.name, a.awarded_by): a
        for a in person.awards
        if a.name and a.awarded_by
    }

    for award_data in awards_data:
        name = award_data.get("name")
        awarded_by = award_data.get("awarded_by")
        if not name or not awarded_by:
            continue

        key = (name, awarded_by)
        start_date, start_precision = normalize_date(award_data.get("start_date"))
        end_date, end_precision = normalize_date(award_data.get("end_date"))

        existing_entry = existing.get(key)
        if existing_entry:
            updated = False
            if (
                existing_entry.start_date != start_date or
                existing_entry.end_date != end_date or
                existing_entry.start_date_precision != start_precision or
                existing_entry.end_date_precision != end_precision
            ):
                existing_entry.start_date = start_date
                existing_entry.start_date_precision = start_precision
                existing_entry.end_date = end_date
                existing_entry.end_date_precision = end_precision
                updated = True

            if updated:
                logger.info(f"🔁 Updated award | Person: {person.full_name} | Key: {key}")
        else:
            session.add(Award(
                person_id=person.id,
                name=name,
                awarded_by=awarded_by,
                start_date=start_date,
                start_date_precision=start_precision,
                end_date=end_date,
                end_date_precision=end_precision
            ))
            logger.info(f"➕ Added new award | Person: {person.full_name} | Key: {key}")



def upsert_publications(session, person, pubs_data):
    existing = {
        (p.title, p.journal): p
        for p in person.publications
        if p.title and p.journal
    }

    for pub_data in pubs_data:
        title = pub_data.get("title")
        journal = pub_data.get("journal")
        if not title or not journal:
            continue

        pub_date, pub_precision = normalize_date(pub_data.get("start_date"))
        authors = pub_data.get("authors") or ""

        existing_entry = existing.get((title, journal))
        if existing_entry:
            updated = False
            if existing_entry.publication_date != pub_date:
                existing_entry.publication_date = pub_date
                updated = True
            if existing_entry.publication_date_precision != pub_precision:
                existing_entry.publication_date_precision = pub_precision
                updated = True
            if authors and existing_entry.authors != authors:
                existing_entry.authors = authors
                updated = True
            if updated:
                logger.info(f"🔁 Updated publication: {title} @ {journal} | Person: {person.full_name}")
        else:
            session.add(Publication(
                person_id=person.id,
                title=title,
                journal=journal,
                authors=authors,
                publication_date=pub_date,
                publication_date_precision=pub_precision
            ))
            logger.info(f"➕ Added new publication: {title} @ {journal} | Person: {person.full_name}")

def upsert_personal_achievements(session, person, achievements_data):
    existing = {
        (a.achievement or "", a.description or ""): a
        for a in person.personal_achievements
    }

    for ach in achievements_data:
        title = ach.get("achievement") or ""
        description = ach.get("description") or ""
        if not title:
            continue

        start_date, start_precision = normalize_date(ach.get("start_date"))
        end_date, end_precision = normalize_date(ach.get("end_date"))

        existing_entry = existing.get((title, description))
        if existing_entry:
            updated = False
            if existing_entry.start_date != start_date:
                existing_entry.start_date = start_date
                updated = True
            if existing_entry.start_date_precision != start_precision:
                existing_entry.start_date_precision = start_precision
                updated = True
            if existing_entry.end_date != end_date:
                existing_entry.end_date = end_date
                updated = True
            if existing_entry.end_date_precision != end_precision:
                existing_entry.end_date_precision = end_precision
                updated = True
            if updated:
                logger.info(f"🔁 Updated personal achievement: {title} | Person: {person.full_name}")
        else:
            session.add(PersonalAchievement(
                person_id=person.id,
                achievement=title,
                description=description,
                start_date=start_date,
                start_date_precision=start_precision,
                end_date=end_date,
                end_date_precision=end_precision
            ))
            logger.info(f"➕ Added new personal achievement: {title} | Person: {person.full_name}")


def upsert_private_milestones(session, person, milestones_data):
    existing = {
        (m.event or "", m.description or ""): m
        for m in person.private_milestones
    }

    for ms in milestones_data:
        event = ms.get("event") or ""
        description = ms.get("description") or ""
        if not event:
            continue

        start_date, start_precision = normalize_date(ms.get("start_date"))
        end_date, end_precision = normalize_date(ms.get("end_date"))

        existing_entry = existing.get((event, description))
        if existing_entry:
            updated = False
            if existing_entry.start_date != start_date:
                existing_entry.start_date = start_date
                updated = True
            if existing_entry.start_date_precision != start_precision:
                existing_entry.start_date_precision = start_precision
                updated = True
            if existing_entry.end_date != end_date:
                existing_entry.end_date = end_date
                updated = True
            if existing_entry.end_date_precision != end_precision:
                existing_entry.end_date_precision = end_precision
                updated = True
            if updated:
                logger.info(f"🔁 Updated private milestone: {event} | Person: {person.full_name}")
        else:
            session.add(PrivateMilestone(
                person_id=person.id,
                event=event,
                description=description,
                start_date=start_date,
                start_date_precision=start_precision,
                end_date=end_date,
                end_date_precision=end_precision
            ))
            logger.info(f"➕ Added new private milestone: {event} | Person: {person.full_name}")


##################################################################################################################
#
# MAIN FUNCTION
#
##################################################################################################################

def parse_and_store(doc_id: int) -> int | None:
    session = SessionLocal()
    try:
        doc = session.get(Document, doc_id)
        if not doc:
            logger.error(f"❌ No Document {doc_id}")
            return

        # DELEGATE parsing to new module
        data, prompt, structured = parse_cv_with_llm(doc.source_filename)
        # audit
        doc.llm_prompt = prompt
        doc.llm_response = structured
        session.flush()

        # upsert person + all sections
        person = get_or_create_person(session, data, doc)
        # loop to always default to []
        mapping = [
            ("education", upsert_educations),
            ("professional_experience", upsert_experiences),
            ("languages", upsert_languages),
            ("further_education", upsert_further_education),
            ("certifications", upsert_certifications),
            ("awards", upsert_awards),
            ("publications", upsert_publications),
            ("personal_achievements", upsert_personal_achievements),
            ("private_milestones", upsert_private_milestones),
        ]
        for key, fn in mapping:
            fn(session, person, data.get(key) or [])

        # finally mark parsed
        doc.status = "parsed"
        session.commit()
        logger.info(f"✅ Finished parsing Document {doc_id} for {person.id}")

        # AUTOMATIC NEXT STEPS
        person_id = person.id
        # 1) generate PDF and register
        generate_cv_pdf(person_id, document_id=doc_id)
        # 2) plot timeline and register
        plot_timeline_and_save(person_id, document_id=doc_id)

        # mark complete
        doc.status = "complete"
        session.commit()
        logger.info(f"🎉 All steps done for Document {doc_id}")
        return person_id

    except Exception as e:
        session.rollback()
        logger.error(f"❌ parse_and_store({doc_id}) failed: {e}")
        raise
    finally:
        session.close()
       